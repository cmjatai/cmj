
import io
import logging

from django.conf import settings
from django.http.response import HttpResponse
from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from django.views.generic.base import TemplateView
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Cm

from cmj.mixins import MultiFormatOutputMixin
from sapl.sessao.models import SessaoPlenaria


class RelatorioPautaSessao(TemplateView):

    template_name = 'base.html'

    def get_context_data(self, **kwargs):
        context = TemplateView.get_context_data(self, **kwargs)

        context['object_list'] = SessaoPlenaria.objects.all()
        return context

    def render_to_response(self, context):

        d = Document()

        for s in d.sections:

            s.page_width = Cm(21)
            s.page_height = Cm(29.7)

            cm1, cm2, cm3, cm4 = Cm(1), Cm(2), Cm(3), Cm(4)
            s.top_margin = cm1
            s.right_margin = cm1
            s.bottom_margin = cm1
            s.left_margin = cm1

            h = s.header
            h.header_distance = 0

            p = h.paragraphs[0]
            p.style = "Heading 2"
            tabs = p.paragraph_format.tab_stops
            tab_stop = tabs.add_tab_stop(Cm(11), WD_TAB_ALIGNMENT.CENTER)

            bh = p.add_run()
            bh.add_picture(
                settings.FRONTEND_BRASAO_PATH['1024'],
                width=Cm(3)
            )

            p.add_run(
                '\t' + "Câmara Municipal de Jataí - GO"  # For center align of text
            ).bold = True
            #th.style = "Heading 2 Char"

            f = s.footer
            p = f.paragraphs[0]
            p.text = 'rodape'

        #d.add_heading('Document Title', 0)

        tzlt = timezone.localtime()
        p = d.add_paragraph(f'Time: {tzlt} A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        d.add_heading('Heading, level 1', level=1)
        d.add_paragraph('Intense quote', style='Intense Quote')

        d.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        d.add_paragraph(
            'first item in ordered list', style='List Number'
        )

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = d.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

        d.add_page_break()

        d.save(
            f'/home/leandro/Downloads/portalcmj_{self.request.resolver_match.url_name}.docx')

        return HttpResponse(f'docx {tzlt}')

        output = io.BytesIO()
        d.save(output)
        output.seek(0)

        response = HttpResponse(output.read(
        ), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response['Content-Disposition'] = f'attachment; filename="portalcmj_{self.request.resolver_match.url_name}.docx"'
        response['Cache-Control'] = 'no-cache'
        response['Pragma'] = 'no-cache'
        response['Expires'] = 0

        output.close()

        return response
